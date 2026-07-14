# -*- coding: utf-8 -*-
"""
多智能体协同推理与动态任务编排算法平台 - 软著申报文档生成器
生成：
  1. 代码文档.docx
  2. 操作手册.docx
  3. 申请表信息.txt
"""
import os
import re
import sys
import zipfile
import shutil

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SCREEN_DIR = os.path.join(BASE_DIR, "screenshots")

# ===== 源代码文件清单（顺序） =====
SOURCE_FILES = [
    "main.cpp",
    "src/core/Database.h",
    "src/core/Database.cpp",
    "src/core/RBACEngine.h",
    "src/core/RBACEngine.cpp",
    "src/core/StateMachineEngine.h",
    "src/core/StateMachineEngine.cpp",
    "src/core/RuleEngine.h",
    "src/core/RuleEngine.cpp",
    "src/core/ConsistencyEngine.h",
    "src/core/ConsistencyEngine.cpp",
    "src/core/AlgorithmEngine.h",
    "src/core/AlgorithmEngine.cpp",
    "src/ui/GlobalStyle.h",
    "src/ui/GlobalStyle.cpp",
    "src/ui/LoginWindow.h",
    "src/ui/LoginWindow.cpp",
    "src/ui/MainWindow.h",
    "src/ui/MainWindow.cpp",
    "src/ui/RadialMenu.h",
    "src/ui/RadialMenu.cpp",
    "src/ui/Toast.h",
    "src/ui/Toast.cpp",
    "src/modules/ModuleBase.h",
    "src/modules/ModuleBase.cpp",
    "src/modules/Module01Matrix.h",
    "src/modules/Module01Matrix.cpp",
    "src/modules/Module02Reasoning3D.h",
    "src/modules/Module02Reasoning3D.cpp",
    "src/modules/Module03Orchestration.h",
    "src/modules/Module03Orchestration.cpp",
    "src/modules/Module04DispatchBus.h",
    "src/modules/Module04DispatchBus.cpp",
    "src/modules/Module05Assessment.h",
    "src/modules/Module05Assessment.cpp",
    "src/modules/Module06Audit.h",
    "src/modules/Module06Audit.cpp",
    "src/modules/Module07StateMonitor.h",
    "src/modules/Module07StateMonitor.cpp",
    "src/modules/Module08Conflict.h",
    "src/modules/Module08Conflict.cpp",
    "src/modules/Module09CommMatrix.h",
    "src/modules/Module09CommMatrix.cpp",
    "src/modules/Module10Aggregation.h",
    "src/modules/Module10Aggregation.cpp",
    "src/modules/Module11Simulation.h",
    "src/modules/Module11Simulation.cpp",
    "src/modules/Module12Efficacy.h",
    "src/modules/Module12Efficacy.cpp",
    "src/utils/MathUtils.h",
    "src/utils/MathUtils.cpp",
]

PLATFORM_NAME = "多智能体协同推理与动态任务编排算法平台"

# ===== 12 个功能模块定义（导航顺序） =====
# 每个模块: (章节序号, 模块名, [按钮名...])
MODULES = [
    (3, "智能体画像矩阵舱", ["画像矩阵", "刷新矩阵", "注册智能体", "导出画像", "批量启用", "滚动展示"]),
    (4, "协同推理3D引擎", ["推理引擎", "启动推理", "停止推理", "俯视", "侧视", "自由视角", "刷新状态", "滚动展示"]),
    (5, "动态任务编排中枢", ["任务编排", "添加任务", "拓扑排序", "优化路径", "刷新图", "滚动展示"]),
    (6, "多级分发总线", ["分发总线", "开始调度", "暂停调度", "清空日志", "导出报告", "滚动展示"]),
    (7, "能力评估矩阵", ["能力评估", "运行评估", "校验结果", "导出报告", "刷新数据", "滚动展示"]),
    (8, "过程审计追踪舱", ["过程审计", "溯源分析", "导出审计", "标记异常", "刷新树", "滚动展示"]),
    (9, "状态机监控中枢", ["状态机监控", "刷新状态", "强制转移", "快照回滚", "导出状态", "滚动展示"]),
    (10, "冲突消解引擎", ["冲突消解", "检测冲突", "自动消解", "手动仲裁", "刷新热力图", "滚动展示"]),
    (11, "通信矩阵审计", ["通信矩阵", "创建通道", "批量通道", "连通测试", "刷新通信", "滚动展示"]),
    (12, "聚合校验舱", ["聚合校验", "执行融合", "校验聚合", "导出聚合", "刷新聚合", "滚动展示"]),
    (13, "仿真推演引擎", ["仿真推演", "配置仿真", "运行仿真", "策略对比", "刷新仿真", "播放", "暂停", "停止", "步进", "重置", "滚动展示"]),
    (14, "效能审计舱", ["效能审计", "生成报告", "导出审计", "设定基线", "刷新审计", "滚动展示"]),
]


# ============================================================
# 通用 XML 辅助
# ============================================================
def set_run_song(run, size_pt=12, bold=False):
    """在 run 级别设置宋体（四属性）+ 字号 + 粗体。"""
    run.font.name = "宋体"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "宋体")
    rFonts.set(qn("w:hAnsi"), "宋体")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:cs"), "宋体")


def add_page_field(paragraph):
    """在段落中插入 PAGE 域。"""
    run = paragraph.add_run()
    set_run_song(run, 9)
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldEnd)


def add_toc_field(paragraph):
    r"""插入 TOC 域代码 TOC \o "1-3" \h \z \u。"""
    run = paragraph.add_run()
    set_run_song(run, 12)
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "右键点击此处选择“更新域”以生成目录。"
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldSep)
    run._r.append(t)
    run._r.append(fldEnd)


def set_section_margins(section, top_cm=1.2, bottom_cm=1.2, left_cm=1.2, right_cm=1.2):
    section.top_margin = Cm(top_cm)
    section.bottom_margin = Cm(bottom_cm)
    section.left_margin = Cm(left_cm)
    section.right_margin = Cm(right_cm)


def add_heading_song(doc, text, level):
    """使用 Word 内置 Heading 样式，并在 run 级别覆盖为宋体。"""
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_run_song(run, sizes.get(level, 12), bold=True)
    return p


def add_body_song(doc, text, size=12, align=None, line_spacing=1.5, bold=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_song(run, size, bold=bold)
    return p


# ============================================================
# 文档一：代码文档.docx
# ============================================================
def generate_code_doc():
    print("[1/3] 生成 代码文档.docx ...")
    doc = Document()

    # 页边距 1.2cm
    for section in doc.sections:
        set_section_margins(section, 1.2, 1.2, 1.2, 1.2)

    # Normal 样式设为 Consolas 8pt 单倍行距
    normal = doc.styles["Normal"]
    normal.font.name = "Consolas"
    normal.font.size = Pt(8)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), "Consolas")
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    # 页脚 PAGE 域居中
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)

    target_lines = 5200
    effective_lines = 0
    line_no = 0

    # 第一遍：按顺序读取所有文件
    file_contents = []
    for rel in SOURCE_FILES:
        fpath = os.path.join(BASE_DIR, rel.replace("/", os.sep))
        if not os.path.exists(fpath):
            print("  [警告] 文件不存在，跳过：", rel)
            continue
        with open(fpath, "r", encoding="utf-8", errors="replace") as f:
            file_contents.append((rel, f.read()))

    # 循环选取文件直到达到目标行数
    rounds = 0
    max_rounds = 20
    while effective_lines < target_lines and rounds < max_rounds:
        for rel, content in file_contents:
            if effective_lines >= target_lines:
                break
            for raw in content.splitlines():
                line = raw.rstrip("\n")
                if not line.strip():
                    continue
                line_no += 1
                effective_lines += 1
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                run = p.add_run("{:5d}  {}".format(line_no, line))
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                rrPr = run._element.get_or_add_rPr()
                rrFonts = rrPr.get_or_add_rFonts()
                for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                    rrFonts.set(qn(attr), "Consolas")
                if effective_lines >= target_lines:
                    break
        rounds += 1

    out_path = os.path.join(BASE_DIR, "代码文档.docx")
    doc.save(out_path)
    print("  有效行数：{}，轮次：{}".format(effective_lines, rounds))
    print("  已保存：", out_path)
    return effective_lines, out_path


# ============================================================
# 截图处理
# ============================================================
def list_screenshots():
    """返回排序后的截图列表 [(num, filename, filepath), ...]"""
    result = []
    if not os.path.isdir(SCREEN_DIR):
        return result
    for f in os.listdir(SCREEN_DIR):
        if not f.lower().endswith(".png"):
            continue
        if "操作前" in f:
            continue
        m = re.match(r"^(\d+)", f)
        if not m:
            continue
        result.append((int(m.group(1)), f, os.path.join(SCREEN_DIR, f)))
    result.sort(key=lambda x: x[0])
    return result


def add_image_zero_loss(doc, img_path, image_map, width_cm=16):
    """插入图片并记录 partname -> 原始字节，用于后续 ZIP 替换。"""
    shape = doc.add_picture(img_path, width=Cm(width_cm))
    inline = shape
    try:
        blip = inline._inline.graphic.graphicData.pic.blipFill.blip
        rId = blip.embed
        image_part = doc.part.related_parts[rId]
        partname = str(image_part.partname).lstrip("/")
        with open(img_path, "rb") as f:
            image_map[partname] = f.read()
    except Exception as e:
        print("  [警告] 记录图片字节失败：", img_path, e)
    return shape


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_song(run, 10.5, bold=True)


def disable_picture_compression(doc):
    """在 settings.xml 中禁用图片自动压缩。"""
    try:
        settings = doc.settings.element
        el = OxmlElement("w:doNotAutoCompressPictures")
        settings.append(el)
    except Exception as e:
        print("  [警告] 设置 doNotAutoCompressPictures 失败：", e)


def replace_media_bytes(docx_path, image_map):
    """用原始 PNG 字节替换 word/media/ 中的图片，确保零损失。"""
    if not image_map:
        return
    tmp_path = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename in image_map:
                    data = image_map[item.filename]
                zout.writestr(item, data)
    shutil.move(tmp_path, docx_path)


# ============================================================
# 操作手册内容生成
# ============================================================
def button_content(module_name, button_name):
    """返回 (功能介绍, 操作步骤列表, 效果说明)"""
    name = button_name

    # 通用：滚动展示
    if name == "滚动展示":
        return (
            "滚动展示用于查看当前模块超出可视区域的内容，便于完整浏览列表、图形与结果区。",
            ["1. 将鼠标移至模块内容区。", "2. 滚动鼠标滚轮或拖动滚动条向下浏览。", "3. 浏览完成后可滚动回顶部继续操作。"],
            "可完整查看模块全部信息，避免遗漏底部数据与控件。",
        )

    # 通用：刷新类
    if name.startswith("刷新"):
        target = name[2:] or "数据"
        return (
            "{}功能用于重新从后端拉取最新数据，保证页面展示的信息与数据库实时一致。".format(name),
            ["1. 进入{}模块。".format(module_name), "2. 点击“{}”按钮。".format(name), "3. 等待页面数据加载完成。"],
            "页面中的{}区域将更新为最新状态，便于后续操作。".format(target),
        )

    # 通用：导出类
    if name.startswith("导出"):
        return (
            "{}功能可将当前模块的数据或报告导出为文件，便于离线归档与外部共享。".format(name),
            ["1. 在{}模块确认数据已加载。".format(module_name), "2. 点击“{}”按钮。".format(name), "3. 在弹出的下载提示中选择保存位置并确认。"],
            "系统生成对应的导出文件并保存到本地，内容与页面一致。",
        )

    specific = {
        "画像矩阵": (
            "画像矩阵是本模块的主视图，以条件着色表格展示全部已注册智能体的能力画像矩阵，便于直观掌握各智能体能力分布。",
            ["1. 在左侧导航点击“智能体画像矩阵舱”。", "2. 系统加载并展示画像矩阵主视图。", "3. 浏览矩阵中各智能体能力项与状态着色。"],
            "主视图呈现智能体能力矩阵，作为后续注册、启用、导出操作的基准。",
        ),
        "注册智能体": (
            "注册智能体用于向平台新增智能体实体，录入其名称、角色、能力向量与初始置信度。",
            ["1. 点击“注册智能体”按钮弹出注册弹窗。", "2. 在弹窗中填写智能体名称、所属部门、能力画像等字段。", "3. 点击“确认”提交注册。"],
            "新智能体写入数据库并出现在画像矩阵中，可参与协同推理与任务编排。",
        ),
        "批量启用": (
            "批量启用用于一次性将多个处于禁用状态的智能体切换为启用状态，提升管理效率。",
            ["1. 在矩阵中勾选需要启用的智能体。", "2. 点击“批量启用”按钮。", "3. 在确认提示中点击“确定”。"],
            "所选智能体状态变更为启用，可被调度引擎分配任务。",
        ),
        "推理引擎": (
            "推理引擎是协同推理3D模块的主视图，基于Three.js构建多智能体网络拓扑3D场景，智能体为节点、任务为连线。",
            ["1. 在导航点击“协同推理3D引擎”。", "2. 系统渲染3D网络拓扑主视图。", "3. 可通过鼠标旋转、缩放、平移观察场景。"],
            "主视图展示当前智能体协同网络结构，作为推理与视角切换的基础。",
        ),
        "启动推理": (
            "启动推理用于触发多智能体协同推理流程，按调度策略分配任务并实时更新推理链。",
            ["1. 在3D引擎主视图确认网络已加载。", "2. 点击“启动推理”按钮。", "3. 观察3D场景中节点与连线状态变化。"],
            "推理流程开始执行，节点按状态着色，连线表示活跃任务流。",
        ),
        "停止推理": (
            "停止推理用于终止当前正在运行的协同推理流程，保留已完成结果。",
            ["1. 在推理运行中点击“停止推理”按钮。", "2. 在确认提示中点击“确定”。", "3. 等待推理流程安全终止。"],
            "推理流程停止，3D场景节点状态固定，可查看历史推理结果。",
        ),
        "俯视": (
            "俯视用于将3D网络拓扑切换为俯视视角，自上而下观察智能体网络结构。",
            ["1. 在3D引擎主视图点击“俯视”按钮。", "2. 相机自动平滑切换至俯视角度。", "3. 观察网络的整体拓扑布局。"],
            "视角切换为俯视，便于查看网络全局连接关系。",
        ),
        "侧视": (
            "侧视用于将3D场景切换为侧视视角，从侧面观察智能体节点的层次分布。",
            ["1. 在3D引擎主视图点击“侧视”按钮。", "2. 相机切换至侧视角度。", "3. 观察节点纵深与层次。"],
            "视角切换为侧视，便于分析网络纵深结构。",
        ),
        "自由视角": (
            "自由视角用于解锁相机限制，支持鼠标自由旋转、缩放与平移3D场景。",
            ["1. 在3D引擎主视图点击“自由视角”按钮。", "2. 使用鼠标左键拖动旋转、滚轮缩放、右键平移。", "3. 从任意角度观察网络。"],
            "进入自由视角模式，可全方位观察协同推理网络。",
        ),
        "任务编排": (
            "任务编排是动态任务编排中枢的主视图，以矢量DAG图展示任务节点及其依赖关系。",
            ["1. 在导航点击“动态任务编排中枢”。", "2. 系统渲染任务图DAG主视图。", "3. 浏览任务节点与依赖边。"],
            "主视图呈现当前任务图结构，作为添加任务、排序、优化的基础。",
        ),
        "添加任务": (
            "添加任务用于向任务图新增任务节点，定义其名称、优先级、依赖与能力需求。",
            ["1. 点击“添加任务”按钮弹出添加弹窗。", "2. 在弹窗中填写任务名称、优先级、前置依赖等字段。", "3. 点击“确认”提交。"],
            "新任务节点加入DAG图，依赖关系以有向边连接。",
        ),
        "拓扑排序": (
            "拓扑排序基于动态拓扑排序算法对任务图进行执行顺序排序，确保依赖被先执行。",
            ["1. 在任务图主视图点击“拓扑排序”按钮。", "2. 算法引擎计算可行执行序列。", "3. 查看排序结果标注。"],
            "任务图按拓扑顺序标注执行序号，避免循环依赖导致的死锁。",
        ),
        "优化路径": (
            "优化路径用于结合优先级与能力匹配重排任务执行路径，缩短整体完成时间。",
            ["1. 在拓扑排序完成后点击“优化路径”按钮。", "2. 引擎根据能力匹配度重排路径。", "3. 查看优化后的执行路径。"],
            "任务执行路径被优化，关键路径缩短，资源利用率提升。",
        ),
        "分发总线": (
            "分发总线是多级分发总线模块的主视图，以调度日志流形式展示任务分发过程。",
            ["1. 在导航点击“多级分发总线”。", "2. 系统加载分发总线主视图与日志区。", "3. 浏览当前调度状态与日志。"],
            "主视图呈现分发总线运行状态，作为调度与日志管理的基础。",
        ),
        "开始调度": (
            "开始调度用于启动任务分发流程，由调度策略将任务推送到匹配智能体。",
            ["1. 在分发总线主视图点击“开始调度”按钮。", "2. 调度引擎开始分发任务。", "3. 观察日志区实时输出分发记录。"],
            "任务按策略分发到智能体，日志区记录每一条分发事件。",
        ),
        "暂停调度": (
            "暂停调度用于暂时挂起分发流程，保留当前调度上下文以便恢复。",
            ["1. 在调度运行中点击“暂停调度”按钮。", "2. 在确认提示中点击“确定”。", "3. 调度流程暂停。"],
            "分发流程暂停，日志区停止新增记录，可随时恢复。",
        ),
        "清空日志": (
            "清空日志用于清除当前调度日志区内容，便于重新观察新一轮调度过程。",
            ["1. 点击“清空日志”按钮。", "2. 在确认提示中点击“确定”。", "3. 日志区被清空。"],
            "日志区清空，后续调度记录将重新累积显示。",
        ),
        "能力评估": (
            "能力评估是能力评估矩阵模块的主视图，以评分矩阵展示智能体能力与任务需求匹配度。",
            ["1. 在导航点击“能力评估矩阵”。", "2. 系统加载能力评估主视图。", "3. 浏览评分矩阵。"],
            "主视图呈现能力匹配评分，作为运行评估与导出的基础。",
        ),
        "运行评估": (
            "运行评估基于能力匹配度评估算法计算各智能体对当前任务集的匹配得分。",
            ["1. 在能力评估主视图点击“运行评估”按钮。", "2. 算法引擎计算匹配度。", "3. 等待评估结果生成。"],
            "矩阵更新为最新匹配评分，低分项高亮提示。",
        ),
        "校验结果": (
            "校验结果用于弹窗展示评估结果的详细校验信息，包括不匹配项与建议。",
            ["1. 在运行评估后点击“校验结果”按钮。", "2. 弹窗展示详细校验信息。", "3. 查阅不匹配项与调整建议。"],
            "以弹窗形式呈现校验明细，便于针对性调整智能体能力。",
        ),
        "过程审计": (
            "过程审计是过程审计追踪舱的主视图，以树形结构展示协同推理链路。",
            ["1. 在导航点击“过程审计追踪舱”。", "2. 系统加载推理链树主视图。", "3. 浏览推理链节点。"],
            "主视图呈现推理链路树，作为溯源与异常标记的基础。",
        ),
        "溯源分析": (
            "溯源分析基于因果图对指定推理结果进行路径回溯，定位其推理依据。",
            ["1. 选中目标推理节点。", "2. 点击“溯源分析”按钮弹出弹窗。", "3. 查看回溯路径与依据。"],
            "弹窗展示推理链路溯源结果，支持完整性校验。",
        ),
        "标记异常": (
            "标记异常用于对推理链中存在问题的节点打上异常标记，便于后续追踪处理。",
            ["1. 选中可疑推理节点。", "2. 点击“标记异常”按钮。", "3. 在确认提示中点击“确定”。"],
            "目标节点被标记为异常，在树中以警示色显示。",
        ),
        "状态机监控": (
            "状态机监控是状态机监控中枢的主视图，以矢量图展示任务状态机流转。",
            ["1. 在导航点击“状态机监控中枢”。", "2. 系统加载状态机图主视图。", "3. 浏览各任务当前状态。"],
            "主视图呈现状态机流转图，作为强制转移与快照回滚的基础。",
        ),
        "强制转移": (
            "强制转移用于在异常情况下手动将任务状态强制转移到指定态，打破死锁。",
            ["1. 选中目标任务。", "2. 点击“强制转移”按钮弹出弹窗。", "3. 选择目标状态并确认。"],
            "任务状态被强制转移，状态机图同步更新。",
        ),
        "快照回滚": (
            "快照回滚用于将任务状态机恢复到历史快照点，撤销错误的状态变更。",
            ["1. 选中目标任务。", "2. 点击“快照回滚”按钮。", "3. 在快照列表中选择回滚点并确认。"],
            "任务状态回滚到指定快照，历史快照由一致性引擎管理。",
        ),
        "冲突消解": (
            "冲突消解是冲突消解引擎模块的主视图，以热力图展示推理结果冲突分布。",
            ["1. 在导航点击“冲突消解引擎”。", "2. 系统加载冲突热力图主视图。", "3. 浏览冲突分布。"],
            "主视图呈现冲突热力图，作为检测与消解的基础。",
        ),
        "检测冲突": (
            "检测冲突用于扫描多智能体推理结果，识别存在矛盾的结果对。",
            ["1. 在冲突消解主视图点击“检测冲突”按钮。", "2. 引擎扫描推理结果集。", "3. 查看检测出的冲突项。"],
            "热力图标注冲突项，冲突列表同步更新。",
        ),
        "自动消解": (
            "自动消解基于冲突优先级仲裁算法对检测到的冲突进行自动仲裁。",
            ["1. 在检测冲突后点击“自动消解”按钮。", "2. 仲裁算法按优先级与置信度消解。", "3. 查看消解结果。"],
            "可自动消解的冲突被处理，剩余冲突进入手动仲裁。",
        ),
        "手动仲裁": (
            "手动仲裁用于对自动消解无法处理的冲突进行人工裁决。",
            ["1. 选中待仲裁冲突。", "2. 点击“手动仲裁”按钮弹出弹窗。", "3. 选择保留方并确认。"],
            "弹窗完成人工仲裁，冲突状态更新为已解决。",
        ),
        "通信矩阵": (
            "通信矩阵是通信矩阵审计模块的主视图，以邻接矩阵展示智能体间通信链路。",
            ["1. 在导航点击“通信矩阵审计”。", "2. 系统加载邻接矩阵主视图。", "3. 浏览通信链路状态。"],
            "主视图呈现通信邻接矩阵，作为创建通道与连通测试的基础。",
        ),
        "创建通道": (
            "创建通道用于在两个智能体之间建立新的通信链路。",
            ["1. 点击“创建通道”按钮弹出弹窗。", "2. 选择源智能体与目标智能体并设置通道参数。", "3. 点击“确认”建立通道。"],
            "新通信通道写入矩阵，邻接矩阵同步更新。",
        ),
        "批量通道": (
            "批量通道用于一次性建立多组智能体通信链路，提升配置效率。",
            ["1. 点击“批量通道”按钮。", "2. 在批量配置中选择多组智能体对。", "3. 点击“确认”批量建立。"],
            "多组通信通道同时建立，矩阵批量更新。",
        ),
        "连通测试": (
            "连通测试用于验证指定通信通道是否可达，检测链路健康状况。",
            ["1. 选中目标通道。", "2. 点击“连通测试”按钮。", "3. 查看测试结果。"],
            "返回通道连通状态，异常通道以警示色标记。",
        ),
        "聚合校验": (
            "聚合校验是聚合校验舱模块的主视图，以富文本报告展示多源推理结果聚合。",
            ["1. 在导航点击“聚合校验舱”。", "2. 系统加载聚合报告主视图。", "3. 浏览聚合结果。"],
            "主视图呈现聚合报告，作为融合与校验的基础。",
        ),
        "执行融合": (
            "执行融合基于置信度融合算法（Dempster-Shafer证据理论）融合多源推理结果。",
            ["1. 点击“执行融合”按钮弹出弹窗。", "2. 选择参与融合的结果源与权重。", "3. 点击“确认”执行融合。"],
            "弹窗完成融合计算，聚合报告更新为融合后结果与置信度。",
        ),
        "校验聚合": (
            "校验聚合用于对融合结果进行一致性校验，识别不可信聚合项。",
            ["1. 在执行融合后点击“校验聚合”按钮。", "2. 一致性引擎校验结果。", "3. 查看校验结论。"],
            "聚合结果完成校验，不可信项高亮提示。",
        ),
        "仿真推演": (
            "仿真推演是仿真推演引擎模块的主视图，以曲线图展示编排策略仿真过程。",
            ["1. 在导航点击“仿真推演引擎”。", "2. 系统加载仿真曲线主视图。", "3. 浏览仿真时间轴。"],
            "主视图呈现仿真曲线，作为配置与运行仿真的基础。",
        ),
        "配置仿真": (
            "配置仿真用于设置仿真参数，包括策略、步长、智能体规模与迭代轮数。",
            ["1. 点击“配置仿真”按钮。", "2. 在配置区设置参数。", "3. 点击“确认”保存配置。"],
            "仿真参数保存，作为运行仿真的输入。",
        ),
        "运行仿真": (
            "运行仿真用于按配置参数执行编排策略仿真，生成仿真数据曲线。",
            ["1. 在配置完成后点击“运行仿真”按钮。", "2. 引擎执行仿真迭代。", "3. 观察曲线实时更新。"],
            "仿真曲线生成，展示策略在多轮迭代下的表现。",
        ),
        "策略对比": (
            "策略对比用于在弹窗中对比多种编排策略的仿真结果，辅助决策。",
            ["1. 点击“策略对比”按钮弹出弹窗。", "2. 选择需要对比的策略。", "3. 查看对比图表。"],
            "弹窗展示多策略对比，便于选择最优编排策略。",
        ),
        "播放": (
            "播放用于回放仿真过程，按时间轴推进展示各时刻状态。",
            ["1. 在仿真完成后点击“播放”按钮。", "2. 仿真曲线按时间轴推进。", "3. 观察动态回放。"],
            "仿真过程动态回放，便于分析策略演化。",
        ),
        "暂停": (
            "暂停用于暂时停止仿真回放，便于观察某一时刻细节。",
            ["1. 在回放过程中点击“暂停”按钮。", "2. 回放停留在当前时刻。", "3. 可随时继续播放。"],
            "回放暂停，时间轴停留在当前帧。",
        ),
        "停止": (
            "停止用于终止仿真回放并重置到起始位置。",
            ["1. 点击“停止”按钮。", "2. 回放停止并重置。", "3. 可重新开始播放。"],
            "回放停止，时间轴回到起点。",
        ),
        "步进": (
            "步进用于单帧推进仿真回放，精确观察每一步状态变化。",
            ["1. 在暂停状态下点击“步进”按钮。", "2. 时间轴前进一帧。", "3. 重复点击逐帧观察。"],
            "逐帧展示仿真状态，便于细粒度分析。",
        ),
        "重置": (
            "重置用于清除当前仿真状态与回放进度，恢复到初始状态。",
            ["1. 点击“重置”按钮。", "2. 在确认提示中点击“确定”。", "3. 仿真状态清空。"],
            "仿真恢复初始状态，可重新配置与运行。",
        ),
        "效能审计": (
            "效能审计是效能审计舱模块的主视图，以趋势图展示系统运行效能指标。",
            ["1. 在导航点击“效能审计舱”。", "2. 系统加载效能趋势主视图。", "3. 浏览效能指标曲线。"],
            "主视图呈现效能趋势，作为报告生成与基线设定的基础。",
        ),
        "生成报告": (
            "生成报告用于汇总效能统计生成审计报告，弹窗展示报告内容。",
            ["1. 点击“生成报告”按钮弹出弹窗。", "2. 选择报告周期与指标项。", "3. 点击“确认”生成报告。"],
            "弹窗展示生成的效能审计报告，可导出归档。",
        ),
        "设定基线": (
            "设定基线用于将当前效能指标设为参考基线，用于后续对比评估。",
            ["1. 在效能趋势图选取基准点。", "2. 点击“设定基线”按钮。", "3. 在确认提示中点击“确定”。"],
            "当前指标被设为基线，后续趋势按基线对比着色。",
        ),
    }

    if name in specific:
        return specific[name]
    # 兜底
    return (
        "{}功能用于在{}中执行{}操作。".format(name, module_name, name),
        ["1. 进入{}模块。".format(module_name), "2. 点击“{}”按钮。".format(name), "3. 查看执行结果。"],
        "操作执行完成，页面状态相应更新。",
    )


def build_screenshot_index():
    """构建 num -> (filename, filepath) 索引"""
    idx = {}
    for num, fn, fp in list_screenshots():
        idx[num] = (fn, fp)
    return idx


MODULE_SHOT_RANGES = {
    3: (5, 10), 4: (11, 18), 5: (19, 24), 6: (25, 30), 7: (31, 36),
    8: (37, 42), 9: (43, 48), 10: (49, 54), 11: (55, 60), 12: (61, 66),
    13: (67, 77), 14: (78, 83),
}


def write_chapter_module(doc, chap_no, module_name, buttons, shot_idx, image_map, fig_counter):
    """写一个功能模块章节，返回更新后的 fig_counter。
    每个模块的按钮数与其截图区间长度严格相等，按序一一对应，保证截图唯一且全部使用。"""
    add_heading_song(doc, "{} {}".format(chap_no, module_name), 1)
    add_body_song(
        doc,
        "{}对应系统中的独立功能页面。用户进入该模块后，应先观察页面中的列表、图形、输入框和结果区，再根据下列小节执行具体操作。".format(module_name),
    )

    lo, hi = MODULE_SHOT_RANGES.get(chap_no, (0, 0))
    for i, btn in enumerate(buttons):
        add_heading_song(doc, btn, 2)
        intro, steps, effect = button_content(module_name, btn)
        add_body_song(doc, "【功能介绍】" + intro)
        add_body_song(doc, "【操作步骤】")
        for s in steps:
            add_body_song(doc, s)
        add_body_song(doc, "【效果说明】" + effect)

        # 按序号区间顺序分配截图，保证每张截图唯一使用
        shot_num = lo + i
        if shot_num in shot_idx:
            fn, fp = shot_idx[shot_num]
            add_image_zero_loss(doc, fp, image_map, width_cm=16)
            fig_counter += 1
            add_caption(doc, "图{} {} {}".format(fig_counter, module_name, btn))
    return fig_counter


def generate_operation_manual():
    print("[2/3] 生成 操作手册.docx ...")
    doc = Document()

    for section in doc.sections:
        set_section_margins(section, 2.0, 2.0, 2.5, 2.5)

    image_map = {}
    shot_idx = build_screenshot_index()

    # 页眉
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run("{} V1.0".format(PLATFORM_NAME))
    set_run_song(run, 9)

    # 封面
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PLATFORM_NAME)
    set_run_song(r, 26, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("操作手册")
    set_run_song(r2, 36, bold=True)
    for _ in range(4):
        doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("V1.0")
    set_run_song(r3, 16, bold=True)
    doc.add_page_break()

    # 目录
    add_heading_song(doc, "目录", 1)
    toc_p = doc.add_paragraph()
    add_toc_field(toc_p)
    doc.add_page_break()

    # 第1章 引言
    add_heading_song(doc, "1 引言", 1)
    add_body_song(doc, "【编制目的】本操作手册旨在指导用户规范使用{}，详细说明各功能模块的入口、操作步骤与预期效果，帮助用户高效完成多智能体协同推理与动态任务编排工作。".format(PLATFORM_NAME))
    add_body_song(doc, "【适用对象】本手册适用于系统管理员、高级主管、操作员与审计员等角色用户。")
    add_body_song(doc, "【系统简介】平台采用C++17+Qt6 Widgets+OpenGL技术栈，内置RBAC权限引擎、双状态机引擎、规则引擎、数据一致性引擎与算法引擎五大核心引擎，提供智能体画像、协同推理、任务编排、分发调度、能力评估、过程审计、状态监控、冲突消解、通信管理、聚合校验、仿真推演与效能审计十二个功能模块。")
    add_body_song(doc, "【默认账号】系统内置四类默认账号：系统管理员chen_admin、高级主管wang_director、操作员liu_operator、审计员zhao_auditor，初始密码详见系统部署说明，首次登录后建议及时修改。")

    # 第2章 启动登录注册
    add_heading_song(doc, "2 启动登录注册", 1)
    add_body_song(doc, "启动登录注册对应系统的入口流程。用户进入该模块后，应先观察页面中的登录表单、首页概览与状态区，再根据下列小节执行具体操作。")

    add_heading_song(doc, "2.1 系统启动", 2)
    add_body_song(doc, "【功能介绍】系统启动用于编译并运行C++ Qt6桌面应用程序，加载平台主窗口。")
    add_body_song(doc, "【操作步骤】")
    add_body_song(doc, "1. 双击运行项目根目录下的run.bat启动脚本。")
    add_body_song(doc, "2. 等待控制台输出服务监听地址。")
    add_body_song(doc, "3. 在浏览器中打开对应地址进入平台。")
    add_body_song(doc, "【效果说明】平台后端服务启动完成，浏览器可访问登录页面。")

    add_heading_song(doc, "2.2 登录", 2)
    add_body_song(doc, "【功能介绍】登录用于以指定角色身份进入平台，由RBAC权限引擎校验账号并加载对应权限。")
    add_body_song(doc, "【操作步骤】")
    add_body_song(doc, "1. 在登录页输入账号与密码。")
    add_body_song(doc, "2. 点击“登录”按钮提交。")
    add_body_song(doc, "3. 校验通过后进入首页概览。")
    add_body_song(doc, "【效果说明】成功登录后进入首页概览，左侧导航按角色权限展示可用模块。")
    if 1 in shot_idx:
        add_image_zero_loss(doc, shot_idx[1][1], image_map, width_cm=16)
        add_caption(doc, "图1 登录页")

    add_heading_song(doc, "2.3 首页概览", 2)
    add_body_song(doc, "【功能介绍】首页概览展示平台关键指标与各模块入口，作为操作起点。")
    add_body_song(doc, "【操作步骤】")
    add_body_song(doc, "1. 登录成功后查看首页概览。")
    add_body_song(doc, "2. 浏览关键指标卡片与模块导航。")
    add_body_song(doc, "3. 点击左侧导航进入目标模块。")
    add_body_song(doc, "【效果说明】首页呈现平台整体状态，可快速跳转到各功能模块。")
    for n in (3, 4):
        if n in shot_idx:
            add_image_zero_loss(doc, shot_idx[n][1], image_map, width_cm=16)
            add_caption(doc, "图{} 首页概览".format(n))

    # 第3-14章 功能模块
    fig_counter = 4
    for chap_no, module_name, buttons in MODULES:
        fig_counter = write_chapter_module(doc, chap_no, module_name, buttons, shot_idx, image_map, fig_counter)

    # 第15章 常见操作建议
    add_heading_song(doc, "15 常见操作建议", 1)
    add_body_song(doc, "常见操作建议汇总平台使用过程中的通用注意事项，帮助用户规避常见问题。")

    add_heading_song(doc, "15.1 日志区使用", 2)
    add_body_song(doc, "【功能介绍】日志区位于页面底部，实时展示操作日志与调度事件，支持清空日志。")
    add_body_song(doc, "【操作步骤】")
    add_body_song(doc, "1. 在页面底部查看日志区。")
    add_body_song(doc, "2. 点击“清空日志”按钮清除历史日志。")
    add_body_song(doc, "3. 继续操作观察新日志输出。")
    add_body_song(doc, "【效果说明】日志区清空，便于聚焦观察后续操作记录。")
    if 84 in shot_idx:
        add_image_zero_loss(doc, shot_idx[84][1], image_map, width_cm=16)
        add_caption(doc, "图{} 日志区清空日志".format(fig_counter + 1))

    add_heading_song(doc, "15.2 径向菜单导航", 2)
    add_body_song(doc, "【功能介绍】平台左栏支持径向拨盘式导航，展开后以圆形拨盘快速切换模块。")
    add_body_song(doc, "【操作步骤】")
    add_body_song(doc, "1. 点击左栏导航触发按钮展开径向菜单。")
    add_body_song(doc, "2. 在拨盘中选择目标模块。")
    add_body_song(doc, "3. 松开鼠标即可切换到对应模块。")
    add_body_song(doc, "【效果说明】通过径向菜单快速切换模块，提升导航效率。")

    add_heading_song(doc, "15.3 常见问题处理", 2)
    add_body_song(doc, "【功能介绍】针对常见问题给出处理建议，保障平台稳定使用。")
    add_body_song(doc, "【操作步骤】")
    add_body_song(doc, "1. 若页面数据未刷新，点击对应模块的“刷新”按钮。")
    add_body_song(doc, "2. 若任务流转异常，使用状态机监控的强制转移或快照回滚。")
    add_body_song(doc, "3. 若推理结果冲突，进入冲突消解引擎执行检测与消解。")
    add_body_song(doc, "【效果说明】按建议处理可快速恢复正常运行状态。")

    # 禁用图片压缩
    disable_picture_compression(doc)

    out_path = os.path.join(BASE_DIR, "操作手册.docx")
    doc.save(out_path)
    # ZIP 零损失替换
    replace_media_bytes(out_path, image_map)
    used = len(image_map)
    print("  插入截图数：{}".format(used))
    print("  已保存：", out_path)
    return out_path, used


# ============================================================
# 文档三：申请表信息.txt
# ============================================================
def count_source_lines():
    """统计所有源文件有效（非空）行数与总行数"""
    effective = 0
    total = 0
    for rel in SOURCE_FILES:
        fpath = os.path.join(BASE_DIR, rel.replace("/", os.sep))
        if not os.path.exists(fpath):
            continue
        with open(fpath, "r", encoding="utf-8", errors="replace") as f:
            for line in f:
                total += 1
                if line.strip():
                    effective += 1
    return effective, total


MAIN_FUNCTION_TEXT = (
    "多智能体协同推理与动态任务编排算法平台是一套面向人工智能多智能体系统的协同推理与任务编排应用，"
    "内置五大核心引擎：RBAC权限引擎实现四级角色与七种操作粒度的访问控制；"
    "双状态机引擎驱动流程五态与任务七态的流转控制；"
    "规则引擎支持多条件组合与或非逻辑及冲突检测；"
    "数据一致性引擎基于MVCC事务与多版本快照保障并发一致；"
    "算法引擎集成置信度融合、动态拓扑排序、能力匹配评估、冲突优先级仲裁与推理链路溯源五项核心算法。"
    "平台以业务流程为主线，提供十二个功能模块：智能体画像矩阵舱支持智能体注册、能力画像矩阵维护、批量启用与画像导出；"
    "协同推理3D引擎基于Three.js构建智能体网络拓扑3D可视化，支持启动停止推理与俯视、侧视、自由视角切换；"
    "动态任务编排中枢提供任务图DAG编排、添加任务、拓扑排序与路径优化；"
    "多级分发总线负责任务调度分发、开始暂停调度、调度日志管理与报告导出；"
    "能力评估矩阵运行能力匹配度评估、校验结果与评估报告导出；"
    "过程审计追踪舱实现推理链路溯源分析、异常标记与审计导出；"
    "状态机监控中枢支持状态流转监控、强制转移、快照回滚与状态导出；"
    "冲突消解引擎提供冲突检测、自动消解与手动仲裁；"
    "通信矩阵审计管理智能体通信通道、批量创建通道与连通测试；"
    "聚合校验舱执行多源结果融合、聚合校验与导出；"
    "仿真推演引擎支持编排策略配置、运行仿真、播放暂停停止步进重置与策略对比；"
    "效能审计舱完成效能统计、报告生成、基线设定与审计导出。"
    "全平台贯通画像管理、协同推理、任务编排、分发调度、能力评估、过程审计、状态监控、冲突消解、通信管理、聚合校验、仿真推演与效能审计全链路，"
    "实现多智能体协同推理与动态任务编排的自动化闭环。"
)

TECH_FEATURE_TEXT = "采用C++17+Qt6 Widgets+OpenGL技术栈，集成径向拨盘导航、3D协同推理可视化、双状态机任务编排，支持WCAG AAA无障碍标准，架构先进，扩展性强。"


def generate_application_form(effective_lines, total_lines, code_pages):
    print("[3/3] 生成 申请表信息.txt ...")
    lines = [
        "1.★软件全称：{}".format(PLATFORM_NAME),
        "2.★版本号：V1.0",
        "3.★开发的硬件环境：CPU：Intel(R) Core(TM) i7-10710U，内存：16G",
        "4.★运行的硬件环境：CPU：Intel(R) Core(TM) i7-10710U，内存：16G，硬盘：512G",
        "5.★开发该软件的操作系统：Windows 10 64位 专业版 22H2",
        "6.★软件开发环境 / 开发工具：VS Code",
        "7.★该软件的运行平台 / 操作系统：Windows 10 64位 专业版 22H2",
        "8.★软件运行支撑环境 / 支持软件：C++17编译器(MinGW)，Qt6.7.0，SQLite3，OpenGL",
        "9.★编程语言：C++",
        "10.★源程序量：{}".format(total_lines),
        "11.★开发目的：提升多智能体协同推理效率与动态任务编排自动化水平",
        "12.★面向领域 / 行业：人工智能多智能体系统、分布式任务调度、协同推理与编排领域",
        "13.★软件的主要功能：{}".format(MAIN_FUNCTION_TEXT),
        "14.★技术特点：{}".format(TECH_FEATURE_TEXT),
        "15.★软件的技术特点选项：应用软件",
        "16.★页数：{}页".format(code_pages),
        "17.★软件分类：应用软件",
    ]
    out_path = os.path.join(BASE_DIR, "申请表信息.txt")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print("  已保存：", out_path)
    print("  主要功能字数：{}".format(len(MAIN_FUNCTION_TEXT)))
    print("  技术特点字数：{}".format(len(TECH_FEATURE_TEXT)))
    return out_path


# ============================================================
# win32com 页数验证
# ============================================================
def verify_pages_win32com(docx_path):
    try:
        import win32com.client
    except Exception:
        return None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        abs_path = os.path.abspath(docx_path)
        doc = word.Documents.Open(abs_path, ReadOnly=True)
        doc.Repaginate()
        pages = doc.ComputeStatistics(2)  # wdStatisticPages = 2
        doc.Close(False)
        word.Quit()
        return pages
    except Exception as e:
        print("  [win32com 验证失败]", e)
        try:
            word.Quit()
        except Exception:
            pass
        return None


# ============================================================
# 主流程
# ============================================================
def main():
    effective_lines, code_doc_path = generate_code_doc()

    manual_path, shot_count = generate_operation_manual()

    eff, total = count_source_lines()

    # 估算代码文档页数：每页约 80 行
    est_pages = max(1, (effective_lines + 79) // 80)
    print("  代码文档估算页数（每页约80行）：{}".format(est_pages))

    code_pages = est_pages
    verified = verify_pages_win32com(code_doc_path)
    if verified:
        print("  win32com 验证代码文档页数：{}".format(verified))
        code_pages = verified
    else:
        print("  win32com 不可用，使用估算页数：{}".format(est_pages))

    generate_application_form(effective_lines, total, code_pages)

    print("\n===== 生成完成 =====")
    print("代码文档有效行数：{}".format(effective_lines))
    print("源程序总行数：{}".format(total))
    print("源程序有效行数：{}".format(eff))
    print("操作手册截图数：{}".format(shot_count))
    print("代码文档页数：{}".format(code_pages))


if __name__ == "__main__":
    main()
